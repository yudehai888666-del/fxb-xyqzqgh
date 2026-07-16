from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from flask import current_app
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet


def _lines(markdown):
    return [line.rstrip() for line in markdown.splitlines()]


def export_planning_docx(student, document):
    destination = Path(current_app.config["GENERATED_DIR"]) / "exports" / f"student-{student.id}" / f"plan-{document['id']}-v{document['version']}.docx"
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    section.top_margin = section.bottom_margin = Inches(0.85)
    section.left_margin = section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(document["title"])
    run.bold, run.font.size, run.font.color.rgb = True, Pt(22), RGBColor(31, 78, 121)
    meta = doc.add_paragraph(f"学生：{student.name}    版本：V{document['version']}    状态：{document['status']}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in _lines(document["content_markdown"]):
        if not line:
            doc.add_paragraph()
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(destination)
    return destination


def export_planning_pdf(student, document):
    destination = Path(current_app.config["GENERATED_DIR"]) / "exports" / f"student-{student.id}" / f"plan-{document['id']}-v{document['version']}.pdf"
    destination.parent.mkdir(parents=True, exist_ok=True)
    font_path = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")
    font_name = "Helvetica"
    if font_path.exists():
        try:
            pdfmetrics.registerFont(TTFont("PlanningChinese", str(font_path), subfontIndex=0))
            font_name = "PlanningChinese"
        except Exception:
            pass
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName=font_name, fontSize=10.5, leading=17, spaceAfter=7)
    h1 = ParagraphStyle("H1CN", parent=body, fontSize=17, leading=22, textColor="#1F4E79", spaceBefore=14, spaceAfter=8)
    h2 = ParagraphStyle("H2CN", parent=body, fontSize=14, leading=19, textColor="#1F4E79", spaceBefore=11, spaceAfter=6)
    title = ParagraphStyle("TitleCN", parent=h1, fontSize=22, leading=28, alignment=1, spaceAfter=10)
    story = [Paragraph(document["title"], title), Paragraph(f"学生：{student.name} &nbsp;&nbsp; 版本：V{document['version']} &nbsp;&nbsp; 状态：{document['status']}", body), Spacer(1, 10)]
    for line in _lines(document["content_markdown"]):
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not line:
            story.append(Spacer(1, 7))
        elif line.startswith("# "):
            story.append(Paragraph(safe[2:], h1))
        elif line.startswith("## "):
            story.append(Paragraph(safe[3:], h2))
        elif line.startswith("### "):
            story.append(Paragraph(safe[4:], h2))
        elif line.startswith("- "):
            story.append(Paragraph("• " + safe[2:], body))
        else:
            story.append(Paragraph(safe, body))
    SimpleDocTemplate(str(destination), pagesize=A4, rightMargin=50, leftMargin=50, topMargin=48, bottomMargin=48, title=document["title"], author="学业规划工作台").build(story)
    return destination
